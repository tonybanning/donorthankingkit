"""Generate the starter files a coordinator fills in.

No fake donors, no simulated data. The donation log is created as a blank
fill-in-the-blank sheet: three labelled boxes the user types over. The letter
template is the org's letter with placeholder tokens.

Run directly to (re)generate the files:
    python samples/make_samples.py
"""

import os

import openpyxl
from docx import Document
from docx.shared import Pt

# The three things the tool needs for each donor.
_HEADER = ["Donor Name", "Donation Amount", "Email"]

# The one example row: three boxes to type over.
_EXAMPLE_ROW = ["Insert name here", "Insert donation here", "Insert email here"]

# Any cell still containing this word means the sheet hasn't been filled in yet.
PLACEHOLDER_MARKER = "insert"


def generate_donations(path):
    """Write the blank starter donation sheet to `path` (.xlsx).

    One header row plus one example row of three boxes. The coordinator replaces
    the boxes with a real donor and adds a row per additional donor.
    """
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Donations"
    sheet.append(_HEADER)
    sheet.append(_EXAMPLE_ROW)

    # Make it obvious and easy to type into.
    for column, width in zip(("A", "B", "C"), (28, 20, 32)):
        sheet.column_dimensions[column].width = width
    for cell in sheet[1]:
        cell.font = openpyxl.styles.Font(bold=True)

    workbook.save(path)


def generate_template(path):
    """Write the customizable letter template to `path` (.docx).

    Laid out like a normal donation letter: a letterhead (logo, organisation
    name, slogan) followed by the body. Every box is editable in Word - reword,
    reorder, restyle, or delete anything. Only the [TOKENS] are filled in for you.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()

    def centered(text, bold=False, italic=False, size=None):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        return paragraph

    # --- Letterhead (same on every letter) ---
    centered("[YOUR LOGO HERE]")           # swapped for your logo.png, if you add one
    centered("[ORGANIZATION]", bold=True, size=18)
    centered("[YOUR SLOGAN HERE]", italic=True)
    document.add_paragraph("")

    # --- Body (personalised per donor) ---
    document.add_paragraph("Dear [NAME],")
    document.add_paragraph("")
    document.add_paragraph(
        "Thank you for your donation of [AMOUNT] to [ORGANIZATION], received [DATE]. "
        "Contributions like yours make our work possible."
    )
    document.add_paragraph("")
    document.add_paragraph("[PERSONAL_NOTE]")
    document.add_paragraph("")
    document.add_paragraph("With gratitude,")
    document.add_paragraph("The Team at [ORGANIZATION]")
    document.save(path)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    generate_donations(os.path.join(here, "donations.xlsx"))
    generate_template(os.path.join(here, "letter_template.docx"))
    print("Wrote donations.xlsx and letter_template.docx - open them and fill them in.")
