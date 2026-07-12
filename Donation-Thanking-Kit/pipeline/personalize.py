"""Stage 2 - Personalize.

For each 'ready' donor: take the org's own Word template, fill the deterministic
placeholders ([NAME], [AMOUNT], [DATE]) in code, and let Claude supply just the
[PERSONAL_NOTE] sentence. The org's wording and structure are never rewritten --
Claude only fills the one placeholder reserved for it.
"""

import os
import re
from datetime import date as date_cls

from docx import Document
from docx.shared import Inches

import config
from pipeline import ai, offline


def _slug(name, index):
    """Filesystem-safe, collision-proof base name for a donor's files."""
    base = re.sub(r"[^a-zA-Z0-9]+", "_", (name or "").strip().lower()).strip("_")
    return f"{index:02d}_{base or 'donor'}"


def _friendly_date(iso):
    """'2026-01-02' -> 'January 2, 2026'; None -> 'recently' (reads naturally)."""
    if not iso:
        return "recently"
    try:
        year, month, day = (int(part) for part in iso.split("-"))
        dt = date_cls(year, month, day)
        return f"{dt.strftime('%B')} {dt.day}, {dt.year}"
    except (ValueError, TypeError):
        return "recently"


def _format_amount(amount):
    try:
        return f"${float(amount):,.2f}"
    except (TypeError, ValueError):
        return str(amount)


def _replace_in_paragraph(paragraph, mapping):
    """Replace placeholders even when Word split them across runs."""
    full = "".join(run.text for run in paragraph.runs)
    if not any(token in full for token in mapping):
        return
    for token, value in mapping.items():
        full = full.replace(token, value)
    if paragraph.runs:
        paragraph.runs[0].text = full
        for run in paragraph.runs[1:]:
            run.text = ""
    else:  # pragma: no cover - paragraphs without runs are rare
        paragraph.text = full


def _fill(document, mapping):
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, mapping)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, mapping)


def _delete_paragraph(paragraph):
    """Remove a paragraph from the document entirely."""
    element = paragraph._element
    element.getparent().remove(element)


def _drop_unfilled_boxes(document, mapping):
    """Delete any paragraph that is *only* a box the org left blank.

    Without this, a template still saying "[YOUR SLOGAN HERE]" would mail that
    text straight to a donor. A blank box is silently dropped instead.
    """
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text in mapping and not str(mapping[text]).strip():
            _delete_paragraph(paragraph)


def _apply_logo(document, logo_path):
    """Swap the [YOUR LOGO HERE] box for the org's actual logo image.

    If they never added a logo file, the box is removed so donors never see it.
    """
    for paragraph in list(document.paragraphs):
        if config.PLACEHOLDER_LOGO not in paragraph.text:
            continue
        if not logo_path:
            _delete_paragraph(paragraph)
            continue
        # Clear the placeholder text, then drop the image into the same spot so
        # it keeps the template's alignment (centred, in the default template).
        for run in paragraph.runs:
            run.text = ""
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.add_picture(logo_path, width=Inches(config.LOGO_WIDTH_INCHES))


def _document_text(document):
    """Plain-text version of the letter, for the email body.

    Blank paragraphs collapse so the email doesn't open with a wall of
    whitespace where the letterhead was.
    """
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(lines)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def run(records, template_path, letters_dir):
    """Write one .docx per ready donor.

    Returns a list of dicts: {record, letter_path, body_text}.
    """
    os.makedirs(letters_dir, exist_ok=True)
    results = []
    ready = [r for r in records if r["status"] == "ready"]

    logo_path = config.find_logo()
    if logo_path:
        print(f"    Using your logo: {os.path.basename(logo_path)}")
    if config.ORG_NAME == "Your Organization Name":
        print(
            "    Tip: set ORG_NAME (and ORG_SLOGAN) in config.py so your letters\n"
            "         carry your organisation's name."
        )

    use_ai = ai.enabled()

    for index, record in enumerate(ready, start=1):
        # Free by default; Claude only if a key is present. If the API call
        # fails we quietly use the free sentence rather than losing the letter.
        note = ai.draft_personal_note(record) if use_ai else None
        if not note:
            note = offline.personal_note(record, index)

        mapping = {
            # Different for every donor:
            config.PLACEHOLDER_NAME: record["donor_name"],
            config.PLACEHOLDER_AMOUNT: _format_amount(record["amount"]),
            config.PLACEHOLDER_DATE: _friendly_date(record.get("date")),
            config.PLACEHOLDER_NOTE: note,
            # Your letterhead, the same on every letter:
            config.PLACEHOLDER_ORG: config.ORG_NAME,
            config.PLACEHOLDER_SLOGAN: config.ORG_SLOGAN,
        }

        # Reload the template fresh each time so donors never bleed into each other.
        document = Document(template_path)
        _apply_logo(document, logo_path)
        _drop_unfilled_boxes(document, mapping)
        _fill(document, mapping)

        letter_path = os.path.join(letters_dir, _slug(record["donor_name"], index) + ".docx")
        document.save(letter_path)
        print(f"    - letter for {record['donor_name']} -> {os.path.basename(letter_path)}")

        results.append({
            "record": record,
            "letter_path": letter_path,
            "body_text": _document_text(document),
        })
    return results
