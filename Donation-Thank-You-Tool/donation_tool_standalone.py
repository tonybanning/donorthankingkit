# ==========================================================================
#  DONATION THANK-YOU TOOL  -  one-file, copy-paste version
# ==========================================================================
#
#  This is the WHOLE tool in a single file. It cleans your donation
#  spreadsheet, writes a thank-you letter for each donor, and prepares a
#  ready-to-send email for each one. It NEVER sends email - you review
#  everything and click Send yourself.
#
#  *** IT IS FREE. ***
#  No API key. No account. No internet. No cost, ever. Everything runs on your
#  own computer and your donor data never leaves it. Just paste and run.
#
#  ------------------------------------------------------------------
#  HOW TO RUN IT - Thonny (free, on your own computer):
#  ------------------------------------------------------------------
#   1. Install Thonny from  https://thonny.org  (it includes Python).
#   2. Open Thonny, paste this whole file in, press the green Run button.
#   3. The first run creates "donations.xlsx" and "letter_template.docx",
#      then stops. Fill in the spreadsheet, then press Run again.
#   4. Read output/REVIEW.md, then open the emails in output/outbox/ and send
#      the ones you approve.
#
#  ------------------------------------------------------------------
#  MAKING THE LETTER YOURS
#  ------------------------------------------------------------------
#   "letter_template.docx" is a normal Word document - reword it, restyle it,
#   rearrange it however you like. It has fill-in boxes for your letterhead:
#
#     [YOUR LOGO HERE]     -> save your logo as "logo.png" beside this file
#     [ORGANIZATION]       -> set ORG_NAME in the SETTINGS section below
#     [YOUR SLOGAN HERE]   -> set ORG_SLOGAN in the SETTINGS section below
#
#   Any box you leave blank is simply removed - a donor will never receive a
#   letter that still says "[YOUR SLOGAN HERE]".
#
#  ------------------------------------------------------------------
#  OPTIONAL: better letters using Claude AI (this part DOES cost money)
#  ------------------------------------------------------------------
#   Leave USE_CLAUDE_AI = False and everything works for free.
#   Set it to True only if you have a paid Anthropic API key and want a
#   freshly-written sentence for each donor. Nothing else changes.
# ==========================================================================

import importlib
import os
import subprocess
import sys


# ==========================================================================
#  SETTINGS  -  edit the text between the quotes to make the letters yours
# ==========================================================================

# Your organisation's name. Fills every [ORGANIZATION] box in the letter.
ORG_NAME = "Your Organization Name"

# Your tagline. Fills the [YOUR SLOGAN HERE] box. Leave it as "" if you don't
# have one - that line is then simply removed from the letter.
ORG_SLOGAN = ""

# Your logo: save an image named logo.png (or logo.jpg) next to this file and it
# is dropped into the [YOUR LOGO HERE] box automatically. No logo? The box is
# removed and the letter still looks fine.
LOGO_FILENAMES = ("logo.png", "logo.jpg", "logo.jpeg")
LOGO_WIDTH_INCHES = 1.5

# Subject line of the prepared emails. {org} and {name} get filled in.
SUBJECT_TEMPLATE = "Thank you for your donation to {org}"

# --- FREE vs PAID ----------------------------------------------------------
# False = FREE. Runs entirely on your computer. No key, no internet, no cost.
# True  = use Claude AI to write a fresher sentence per donor. Needs a paid
#         Anthropic API key, which it will ask you for. Everything else is
#         identical - this ONLY changes the wording of the personal sentence
#         and how messy a spreadsheet it can untangle.
USE_CLAUDE_AI = False

# Only used when USE_CLAUDE_AI is True.
# Model IDs: https://docs.claude.com/en/docs/about-claude/models/overview
CLEAN_MODEL = "claude-haiku-4-5"
DRAFT_MODEL = "claude-sonnet-5"

# --- Boxes you can put in the Word letter ---------------------------------
# Filled differently for every donor:
PLACEHOLDER_NAME = "[NAME]"
PLACEHOLDER_AMOUNT = "[AMOUNT]"
PLACEHOLDER_DATE = "[DATE]"
PLACEHOLDER_NOTE = "[PERSONAL_NOTE]"
# The same on every letter (your letterhead):
PLACEHOLDER_ORG = "[ORGANIZATION]"
PLACEHOLDER_SLOGAN = "[YOUR SLOGAN HERE]"
PLACEHOLDER_LOGO = "[YOUR LOGO HERE]"

# Where results are written.
OUTPUT_DIR = "output"
LETTERS_DIR = os.path.join(OUTPUT_DIR, "letters")
OUTBOX_DIR = os.path.join(OUTPUT_DIR, "outbox")
CLEANED_CSV = os.path.join(OUTPUT_DIR, "cleaned.csv")
REVIEW_MD = os.path.join(OUTPUT_DIR, "REVIEW.md")
SEND_HTML = os.path.join(OUTPUT_DIR, "SEND.html")

# Open SEND.html in your browser automatically when the run finishes.
OPEN_SEND_PAGE = True


# ==========================================================================
#  Step 0: install the helper libraries (first run only)
# ==========================================================================
def _ensure(pip_name, import_name):
    try:
        importlib.import_module(import_name)
    except ImportError:
        print(f"Installing {pip_name} (one-time setup)...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", pip_name])


_ensure("openpyxl", "openpyxl")       # reads/writes the spreadsheet
_ensure("python-docx", "docx")        # reads/writes the Word letters
# NOTE: the "anthropic" package is only installed if you turn on USE_CLAUDE_AI.

import base64
import csv
import html as html_lib
import json
import mimetypes
import re
import webbrowser
import zipfile
from datetime import date as _date
from datetime import datetime
from email.generator import BytesGenerator
from email.message import EmailMessage
from urllib.parse import quote

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

_ISO_DATE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
_SIMPLE_EMAIL = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
_client = None


def find_logo():
    """Return the path to your logo image, or None if you didn't add one."""
    for filename in LOGO_FILENAMES:
        if os.path.exists(filename):
            return filename
    return None


# ==========================================================================
#  FREE MODE  -  everything the AI would do, done in plain Python
# ==========================================================================

_WORD_NUMBERS = {
    "zero": 0, "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6,
    "seven": 7, "eight": 8, "nine": 9, "ten": 10, "eleven": 11, "twelve": 12,
    "thirteen": 13, "fourteen": 14, "fifteen": 15, "sixteen": 16,
    "seventeen": 17, "eighteen": 18, "nineteen": 19, "twenty": 20,
    "thirty": 30, "forty": 40, "fifty": 50, "sixty": 60, "seventy": 70,
    "eighty": 80, "ninety": 90,
}

_DATE_FORMATS = (
    "%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y", "%d/%m/%Y", "%Y/%m/%d",
    "%B %d, %Y", "%B %d %Y", "%b %d, %Y", "%b %d %Y", "%d %B %Y", "%m-%d-%Y",
)


def parse_amount(text):
    """'$1,250.00' / '50' / 'fifty' -> a number. None if it isn't one."""
    if text is None or str(text).strip() == "":
        return None
    if isinstance(text, (int, float)):
        return float(text)

    stripped = re.sub(r"[^0-9.\-]", "", str(text))
    if stripped not in ("", "-", ".", "-."):
        try:
            return float(stripped)
        except ValueError:
            pass

    words = re.findall(r"[a-z]+", str(text).lower())
    if not words:
        return None
    total, current, seen = 0, 0, False
    for word in words:
        if word in _WORD_NUMBERS:
            current += _WORD_NUMBERS[word]
            seen = True
        elif word == "hundred" and seen:
            current = (current or 1) * 100
        elif word == "thousand" and seen:
            total += (current or 1) * 1000
            current = 0
        elif word in ("and", "dollars", "dollar", "usd", "bucks"):
            continue
        else:
            return None  # a word we don't understand -> don't guess
    return float(total + current) if seen else None


def parse_date(text):
    """Any common date format -> 'YYYY-MM-DD'. None if unreadable."""
    text = str(text or "").strip()
    if not text:
        return None
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(text, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return None


def tidy_name(text):
    """'  SMITH, ROBERT ' -> 'Robert Smith'.  'bob smith' -> 'Bob Smith'."""
    name = " ".join(str(text or "").split())
    if not name:
        return ""
    if "," in name:  # "Last, First" -> "First Last"
        last, _, first = name.partition(",")
        name = f"{first.strip()} {last.strip()}".strip()
    if name.isupper() or name.islower():  # leave "McDonald" alone
        name = name.title()
    return name


def _find_column(header, keywords):
    for index, cell in enumerate(header):
        lowered = str(cell).lower()
        if any(keyword in lowered for keyword in keywords):
            return index
    return None


def _cell(row, index):
    if index is None or index >= len(row):
        return ""
    return str(row[index]).strip()


def clean_rows_free(header, rows):
    """Tidy the spreadsheet with no AI: parse, normalise, merge, flag."""
    name_at = _find_column(header, ("name", "donor", "supporter"))
    amount_at = _find_column(header, ("amount", "donation", "gift", "sum", "total", "$"))
    email_at = _find_column(header, ("email", "e-mail", "mail"))
    date_at = _find_column(header, ("date", "received", "when"))

    if name_at is None:
        name_at = 0
    if email_at is None and rows:  # headers renamed? find the column with the @s
        for index in range(len(rows[0])):
            if any("@" in _cell(row, index) for row in rows):
                email_at = index
                break

    dates_fixed = 0
    gifts = []
    for row in rows:
        raw_amount, raw_date = _cell(row, amount_at), _cell(row, date_at)
        amount, date = parse_amount(raw_amount), parse_date(raw_date)
        if raw_date and date and raw_date != date:
            dates_fixed += 1

        notes = []
        if raw_amount and amount is None:
            notes.append(f"could not read the donation amount '{raw_amount}'")
        if raw_date and date is None:
            notes.append(f"could not read the date '{raw_date}'")

        gifts.append({
            "donor_name": tidy_name(_cell(row, name_at)),
            "email": _cell(row, email_at).lower() or None,
            "amount": amount,
            "date": date,
            "notes": notes,
        })

    # Group into one record per donor (same email = same person).
    groups, order = {}, []
    for gift in gifts:
        key = gift["email"] or f"name:{gift['donor_name'].lower()}"
        if key not in groups:
            groups[key] = []
            order.append(key)
        groups[key].append(gift)

    merged, records = 0, []
    for key in order:
        group = groups[key]
        if len(group) > 1:
            merged += len(group) - 1
        latest = sorted(group, key=lambda g: g["date"] or "")[-1]
        amounts = [g["amount"] for g in group if g["amount"] is not None]

        notes = []
        for gift in group:
            notes.extend(gift["notes"])
        if len(group) > 1:
            notes.insert(
                0, f"recurring donor: {len(group)} gifts totalling ${sum(amounts):,.2f}"
            )

        records.append({
            "donor_name": latest["donor_name"],
            "email": latest["email"],
            "amount": latest["amount"],
            "date": latest["date"],
            "donor_type": "recurring" if len(group) > 1 else "first-time",
            "status": "ready",  # re-checked below; anything incomplete gets flagged
            "notes": "; ".join(notes),
        })

    change_log = [f"read {len(gifts)} row(s) from the spreadsheet"]
    if merged:
        change_log.append(f"merged {merged} duplicate row(s) into existing donors")
    if dates_fixed:
        change_log.append(f"normalised {dates_fixed} date(s)")
    change_log.append("cleaned offline - free, and no data left your computer")
    return {"records": records, "change_log": change_log}


# The personal sentence, without an AI.
#
# The tone climbs with the gift. A small gift gets quiet, humble, genuine
# gratitude - never gushing, because overselling a $10 gift feels false. A
# transformational gift gets to hear what it truly did: that a life was changed,
# and that the donor is the reason. Four sizes, two donor types, no form letters.
_NOTES = {
    # --- Under $25: humble, sincere, human. No fireworks. -------------------
    ("first-time", "small"): [
        "It means a great deal that you chose to give at all, and we don't take that lightly.",
        "Thank you for taking that first step with us. Gifts this size keep the lights on, quietly and reliably.",
        "We're simply glad you're here. Every gift counts, and yours is no exception.",
    ],
    ("recurring", "small"): [
        "You keep coming back, and we notice. That steadiness is worth more to us than any single large gift.",
        "Thank you for giving again. Faithful friends like you are the quiet backbone of everything we manage to do.",
        "Your support arrives without fail, and we've come to depend on it more than you probably realize.",
    ],
    # --- $25-$249: a real contribution. Warm and specific. ------------------
    ("first-time", "medium"): [
        "Your first gift is a real encouragement, and we're grateful you decided to trust us with it.",
        "Thank you for choosing to support this work. A first gift of this size opens real doors for us.",
        "You've arrived at exactly the right moment, and we're glad to welcome you into this community.",
    ],
    ("recurring", "medium"): [
        "Your continued generosity is something we genuinely plan around, and we don't take it for granted.",
        "Year after year, you keep showing up for this work. Thank you for being someone we can count on.",
        "Supporters who stay the course are what make long-term work possible - and you have stayed the course.",
    ],
    # --- $250-$999: significant. Say what it unlocks. -----------------------
    ("first-time", "large"): [
        "We're genuinely moved that your very first gift was such a generous one, and we will put every dollar of it to work with care.",
        "A first gift of this size is a serious vote of confidence, and we fully intend to live up to it.",
        "Thank you for beginning your support so generously. This will make a difference people will actually feel.",
    ],
    ("recurring", "large"): [
        "Your sustained and substantial support has shaped what this organization is able to do, and we are deeply grateful to have you in our corner.",
        "Few people give as faithfully, or as generously, as you have. Thank you for continuing to believe in this work.",
        "Your ongoing commitment is felt in everything we accomplish. We would not be who we are without supporters like you.",
    ],
    # --- $1,000+: transformational. Tell them the truth: they changed a life.
    ("first-time", "major"): [
        "It is hard to overstate what a gift like this does. Somewhere out there is a person whose life will be permanently different because you decided to give - and you did it without ever having given to us before. Thank you. Truly.",
        "A gift of this magnitude changes lives, and we mean that literally, not as a figure of speech. You have altered someone's future for good, and we are honored that you chose us to make it happen.",
        "You have done something extraordinary. This gift will reach a real person, at a real moment of need, and change the direction of their life - and you did it on your very first gift. We will never forget it.",
    ],
    ("recurring", "major"): [
        "Year after year you have given at a level that changes lives, and lives have in fact been changed - permanently, and for the better. You are one of the reasons this organization exists at all. Thank you, from all of us.",
        "There are people walking a different path today because of you. Your continued generosity at this level is not simply support; it is the thing itself, the reason the work is possible. We are profoundly grateful.",
        "Your faithfulness at this magnitude has transformed lives, and it has transformed us. We do not have adequate words - only the certainty that what you have made possible will outlast all of us.",
    ],
}

_FALLBACK_NOTE = (
    "Your generosity means a great deal to us, and we are grateful for your support."
)


def _tier(amount):
    """Gift size band. Drives how warm the personal sentence gets.

    The humble register is reserved for genuinely token gifts (under $25).
    Anything from $25 up is a real, considered contribution and is thanked
    warmly - underplaying it would be its own kind of insult.
    """
    try:
        amount = float(amount)
    except (TypeError, ValueError):
        return "medium"
    if amount < 25:
        return "small"       # token gift - quiet, humble thanks
    if amount < 250:
        return "medium"      # a real contribution - warm thanks
    if amount < 1000:
        return "large"       # significant - name what it unlocks
    return "major"           # transformational - you changed a life


def personal_note_free(record, index=0):
    """A warm sentence suited to this donor. Free, instant, offline."""
    donor_type = record.get("donor_type") or "first-time"
    if donor_type not in ("first-time", "recurring"):
        donor_type = "first-time"
    options = _NOTES.get((donor_type, _tier(record.get("amount"))))
    if not options:
        return _FALLBACK_NOTE
    return options[index % len(options)]  # vary it so a batch isn't a mail merge


# ==========================================================================
#  OPTIONAL PAID MODE  -  only runs if USE_CLAUDE_AI = True
# ==========================================================================
def _get_client():
    global _client
    if _client is None:
        _ensure("anthropic", "anthropic")  # installed only if you opted in
        from anthropic import Anthropic

        key = os.environ.get("ANTHROPIC_API_KEY")
        if not key:
            try:
                import getpass
                key = getpass.getpass(
                    "Paste your Anthropic API key (it stays hidden), then Enter: "
                ).strip()
            except Exception:
                key = input("Paste your Anthropic API key: ").strip()
        if not key:
            raise SystemExit(
                "No API key given. Set USE_CLAUDE_AI = False to run for free instead."
            )
        os.environ["ANTHROPIC_API_KEY"] = key
        _client = Anthropic(api_key=key)
    return _client


def _text_of(response):
    return "".join(b.text for b in response.content if b.type == "text")


def _extract_json(text):
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text.strip())
    start, end = text.find("{"), text.rfind("}")
    if start != -1 and end != -1 and end > start:
        text = text[start:end + 1]
    return json.loads(text)


def clean_rows_ai(header, rows):
    payload = {"columns": header, "rows": rows}
    instructions = f"""\
Here is a raw donation log as JSON (column headers plus rows):

{json.dumps(payload, ensure_ascii=False)}

Produce ONE record per unique donor using this exact schema:

  donor_name : full name, tidy capitalisation (e.g. "Robert Smith")
  email      : lowercase email, or null if genuinely absent
  amount     : the gift amount as a NUMBER (convert "$50", "50.00", "fifty" -> 50).
               If a donor gave more than once, use their MOST RECENT gift here.
  date       : most recent gift date as "YYYY-MM-DD", or null if absent/unparseable
  donor_type : "recurring" if the donor appears more than once, else "first-time"
  status     : "ready" if donor_name, email and amount are all present;
               "flagged" if any required field is missing; "skipped" for junk rows
  notes      : short human-readable explanation of anything you changed or flagged

Rules:
  - Merge rows that are obviously the same donor (same email, or the same name in
    different formats like "Bob Smith" / "Robert Smith"). Never merge different people.
  - Never fabricate a missing email or amount. Flag instead.

Reply with JSON only:
{{"records": [...], "change_log": ["short sentence", ...]}}
"""
    resp = _get_client().messages.create(
        model=CLEAN_MODEL,
        max_tokens=8000,
        system=(
            "You clean messy donation logs for a small nonprofit. You never "
            "invent data. You reply with JSON only."
        ),
        messages=[{"role": "user", "content": instructions}],
    )
    return _extract_json(_text_of(resp))


def personal_note_ai(record):
    """One tailored sentence from Claude, or None if the call fails."""
    prompt = (
        f"Donor: {record.get('donor_name')}\n"
        f"Relationship: {record.get('donor_type') or 'first-time'}\n"
        f"Most recent gift: {record.get('amount')}\n"
        f"Notes: {record.get('notes') or ''}\n\n"
        "Write the personalised sentence now."
    )
    try:
        resp = _get_client().messages.create(
            model=DRAFT_MODEL,
            max_tokens=256,
            thinking={"type": "disabled"},  # save the budget for the sentence
            system=(
                "You write the personal note in a nonprofit thank-you letter. "
                "Scale the emotional register to the gift:\n"
                "- Under $25: ONE sentence. Humble, quiet, sincere gratitude. Never "
                "gush - overselling a token gift reads as insincere.\n"
                "- $25-$249: one or two sentences. Warm and specific. This is a real, "
                "considered contribution, not spare change - do not underplay it.\n"
                "- $250-$999: two sentences. Significant. Name what it makes possible.\n"
                "- $1,000+: two to four sentences. Transformational. Tell them "
                "plainly that they changed a life - that a real person's future is "
                "permanently different because of them - and mean it. This is the "
                "one case where real emotional weight is warranted.\n"
                "Recurring donors should hear that their faithfulness is noticed. "
                "Never restate the dollar amount mechanically, never make promises "
                "you can't keep, never add a greeting or sign-off. Output only the "
                "note itself, as plain text."
            ),
            messages=[{"role": "user", "content": prompt}],
        )
        return _text_of(resp).strip() or None
    except Exception as exc:
        print(f"    ! could not tailor a note for {record.get('donor_name')}: {exc}")
        return None


# ==========================================================================
#  Starter files (blank templates to fill in - no fake donors)
# ==========================================================================
def generate_donations(path):
    """Create a blank donation sheet: three labelled boxes to type over."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Donations"
    ws.append(["Donor Name", "Donation Amount", "Email"])
    ws.append(["Insert name here", "Insert donation here", "Insert email here"])
    for column, width in zip(("A", "B", "C"), (28, 20, 32)):
        ws.column_dimensions[column].width = width
    for cell in ws[1]:
        cell.font = openpyxl.styles.Font(bold=True)
    wb.save(path)


def generate_template(path):
    """Create the customizable letter, laid out like a normal donation letter."""
    doc = Document()

    def centered(text, bold=False, italic=False, size=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        return p

    # --- Letterhead (the same on every letter) ---
    centered("[YOUR LOGO HERE]")                      # becomes your logo.png
    centered("[ORGANIZATION]", bold=True, size=18)    # becomes ORG_NAME
    centered("[YOUR SLOGAN HERE]", italic=True)       # becomes ORG_SLOGAN
    doc.add_paragraph("")

    # --- Body (personalised for each donor) ---
    doc.add_paragraph("Dear [NAME],")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Thank you for your donation of [AMOUNT] to [ORGANIZATION], received [DATE]. "
        "Contributions like yours make our work possible."
    )
    doc.add_paragraph("")
    doc.add_paragraph("[PERSONAL_NOTE]")
    doc.add_paragraph("")
    doc.add_paragraph("With gratitude,")
    doc.add_paragraph("The Team at [ORGANIZATION]")
    doc.save(path)


# ==========================================================================
#  Stage 1 - Clean
# ==========================================================================
def read_xlsx(path):
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheet = wb.active
    rows = []
    for r in sheet.iter_rows(values_only=True):
        cells = ["" if c is None else str(c).strip() for c in r]
        if any(cells):  # drop the empty trailing rows Excel leaves behind
            rows.append(cells)
    wb.close()
    if not rows:
        return [], []
    return rows[0], rows[1:]


def looks_unfilled(rows):
    """True if the sheet still only holds the 'Insert ... here' boxes."""
    if not rows:
        return True
    for row in rows:
        for cell in row:
            text = str(cell).strip().lower()
            if text and not text.startswith("insert"):
                return False
    return True


def _coerce_amount(value):
    return parse_amount(value)


def _validate(record):
    """Enforce the required fields ourselves, whatever produced the record.

    This is the fail-safe: a donor only stays 'ready' if they genuinely have a
    name, a valid-looking email and an amount.
    """
    notes = []
    if (record.get("notes") or "").strip():
        notes.append(record["notes"].strip())

    record["donor_name"] = (record.get("donor_name") or "").strip()
    email = (record.get("email") or "").strip().lower()
    record["email"] = email or None
    record["amount"] = _coerce_amount(record.get("amount"))

    date = (record.get("date") or "").strip()
    if date and not _ISO_DATE.match(date):
        notes.append(f"could not read the date '{date}'")
        date = ""
    record["date"] = date or None
    record["donor_type"] = record.get("donor_type") or "first-time"

    if record.get("status") == "skipped":
        record["notes"] = "; ".join(n for n in notes if n)
        return record

    missing = []
    if not record["donor_name"]:
        missing.append("name")
    if not record["email"] or not _SIMPLE_EMAIL.match(record["email"]):
        missing.append("email")
    if record["amount"] is None:
        missing.append("donation amount")

    if missing:
        record["status"] = "flagged"
        notes.append("held back - missing/invalid " + ", ".join(missing))
    elif record.get("status") not in ("ready", "flagged", "skipped"):
        record["status"] = "ready"

    record["notes"] = "; ".join(n for n in notes if n)
    return record


def _dedupe_ready(records, change_log):
    seen, result = {}, []
    for r in records:
        if r["status"] == "ready" and r["email"]:
            if r["email"] in seen:
                change_log.append(f"merged duplicate record for {r['email']}")
                continue
            seen[r["email"]] = True
        result.append(r)
    return result


def stage_clean(input_path):
    header, rows = read_xlsx(input_path)
    if not rows:
        return [], ["the spreadsheet had no data rows"]

    if USE_CLAUDE_AI:
        try:
            data = clean_rows_ai(header, rows)
        except Exception as exc:
            print(f"  ! Claude cleaning failed ({exc}); using free mode instead.")
            data = clean_rows_free(header, rows)
    else:
        data = clean_rows_free(header, rows)

    raw = data.get("records", []) if isinstance(data, dict) else []
    change_log = list(data.get("change_log", [])) if isinstance(data, dict) else []
    records = [_validate(dict(r)) for r in raw if isinstance(r, dict)]
    records = _dedupe_ready(records, change_log)

    ready = sum(1 for r in records if r["status"] == "ready")
    flagged = sum(1 for r in records if r["status"] == "flagged")
    skipped = sum(1 for r in records if r["status"] == "skipped")
    change_log.append(f"checked: {ready} ready, {flagged} flagged, {skipped} skipped")
    return records, change_log


# ==========================================================================
#  Stage 2 - Personalize
# ==========================================================================
def _slug(name, index):
    base = re.sub(r"[^a-zA-Z0-9]+", "_", (name or "").strip().lower()).strip("_")
    return f"{index:02d}_{base or 'donor'}"


def _friendly_date(iso):
    if not iso:
        return "recently"
    try:
        y, m, d = (int(p) for p in iso.split("-"))
        dt = _date(y, m, d)
        return f"{dt.strftime('%B')} {dt.day}, {dt.year}"
    except (ValueError, TypeError):
        return "recently"


def _format_amount(amount):
    try:
        return f"${float(amount):,.2f}"
    except (TypeError, ValueError):
        return str(amount)


def _replace_in_paragraph(paragraph, mapping):
    """Replace boxes even when Word split them across several runs."""
    full = "".join(run.text for run in paragraph.runs)
    if not any(tok in full for tok in mapping):
        return
    for tok, val in mapping.items():
        full = full.replace(tok, str(val))
    if paragraph.runs:
        paragraph.runs[0].text = full
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = full


def _fill(document, mapping):
    for p in document.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def _delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def _drop_unfilled_boxes(document, mapping):
    """Delete any line that is only a box you left blank.

    Stops a letter still saying "[YOUR SLOGAN HERE]" from reaching a donor.
    """
    for p in list(document.paragraphs):
        text = p.text.strip()
        if text in mapping and not str(mapping[text]).strip():
            _delete_paragraph(p)


def _apply_logo(document, logo_path):
    """Swap the [YOUR LOGO HERE] box for your actual logo image (or remove it)."""
    for p in list(document.paragraphs):
        if PLACEHOLDER_LOGO not in p.text:
            continue
        if not logo_path:
            _delete_paragraph(p)
            continue
        for run in p.runs:
            run.text = ""
        run = p.runs[0] if p.runs else p.add_run()
        run.add_picture(logo_path, width=Inches(LOGO_WIDTH_INCHES))


def _document_text(document):
    """Plain-text letter, used as the fallback for mail apps that refuse HTML."""
    text = "\n".join(p.text.strip() for p in document.paragraphs)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


# --- Turning the Word letter into a real, formatted email ------------------
#
# An email can't contain a .docx. So we read the finished Word letter back and
# rebuild it as HTML, carrying over the styling the organisation chose in Word:
# bold, italics, underline, font, size, colour, centring - and the logo.

_ALIGNMENTS = {0: "left", 1: "center", 2: "right", 3: "justify"}


def _logo_data_uri(logo_path):
    """Embed the logo directly in the HTML so it travels with the email."""
    if not logo_path or not os.path.exists(logo_path):
        return None
    mime = mimetypes.guess_type(logo_path)[0] or "image/png"
    with open(logo_path, "rb") as handle:
        encoded = base64.b64encode(handle.read()).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def _paragraph_has_image(paragraph):
    return bool(paragraph._element.findall(".//" + qn("w:drawing")))


def _run_to_html(run):
    """One run of text, carrying its Word formatting across to CSS."""
    text = html_lib.escape(run.text)
    if not text:
        return ""
    text = text.replace("\n", "<br>")

    styles = []
    if run.bold:
        styles.append("font-weight:bold")
    if run.italic:
        styles.append("font-style:italic")
    if run.underline:
        styles.append("text-decoration:underline")
    try:
        if run.font.size is not None:
            styles.append(f"font-size:{run.font.size.pt:.0f}pt")
    except (AttributeError, ValueError):
        pass
    try:
        if run.font.color is not None and run.font.color.rgb is not None:
            styles.append(f"color:#{run.font.color.rgb}")
    except (AttributeError, ValueError):
        pass  # theme colours have no .rgb - just leave the default
    if run.font.name:
        styles.append(f"font-family:'{run.font.name}',sans-serif")

    if styles:
        return f'<span style="{";".join(styles)}">{text}</span>'
    return text


def _document_html(document, image_src=None):
    """Rebuild the finished letter as email-ready HTML (inline styles only).

    `image_src` is what the logo <img> points at: a data: URI for the browser
    page, or a cid: reference for the .eml file.
    """
    blocks = []
    for paragraph in document.paragraphs:
        if _paragraph_has_image(paragraph):
            if image_src:
                blocks.append(
                    f'<p style="text-align:center;margin:0 0 10px">'
                    f'<img src="{image_src}" alt="" '
                    f'style="max-width:200px;height:auto"></p>'
                )
            continue

        inner = "".join(_run_to_html(run) for run in paragraph.runs)
        if not inner.strip():
            blocks.append('<p style="margin:0 0 10px">&nbsp;</p>')
            continue

        align = _ALIGNMENTS.get(
            int(paragraph.alignment) if paragraph.alignment is not None else 0, "left"
        )
        blocks.append(
            f'<p style="margin:0 0 10px;text-align:{align}">{inner}</p>'
        )

    return (
        '<div style="font-family:Calibri,Arial,sans-serif;font-size:11pt;'
        'color:#202124;line-height:1.5">' + "".join(blocks) + "</div>"
    )


def stage_personalize(records, template_path):
    os.makedirs(LETTERS_DIR, exist_ok=True)
    results = []
    ready = [r for r in records if r["status"] == "ready"]

    logo_path = find_logo()
    logo_uri = _logo_data_uri(logo_path)
    if logo_path:
        print(f"    Using your logo: {logo_path}")
    if ORG_NAME == "Your Organization Name":
        print("    Tip: set ORG_NAME near the top of this file to your org's name.")

    for i, record in enumerate(ready, start=1):
        note = personal_note_ai(record) if USE_CLAUDE_AI else None
        if not note:
            note = personal_note_free(record, i)

        mapping = {
            # different for every donor
            PLACEHOLDER_NAME: record["donor_name"],
            PLACEHOLDER_AMOUNT: _format_amount(record["amount"]),
            PLACEHOLDER_DATE: _friendly_date(record.get("date")),
            PLACEHOLDER_NOTE: note,
            # your letterhead, the same on every letter
            PLACEHOLDER_ORG: ORG_NAME,
            PLACEHOLDER_SLOGAN: ORG_SLOGAN,
        }

        doc = Document(template_path)
        _apply_logo(doc, logo_path)
        _drop_unfilled_boxes(doc, mapping)
        _fill(doc, mapping)

        letter_path = os.path.join(LETTERS_DIR, _slug(record["donor_name"], i) + ".docx")
        doc.save(letter_path)
        print(f"    - letter for {record['donor_name']}")
        results.append({
            "record": record,
            "letter_path": letter_path,
            "logo_path": logo_path,
            "body_text": _document_text(doc),          # plain-text fallback
            "html_web": _document_html(doc, logo_uri),  # for the browser/clipboard
            "html_mail": _document_html(doc, "cid:logoimg" if logo_path else None),
        })
    return results


# ==========================================================================
#  Stage 3 - Prepare (a human sends)
# ==========================================================================
_DOCX_SUBTYPE = "vnd.openxmlformats-officedocument.wordprocessingml.document"


def _write_eml(letter, eml_path):
    """A real formatted email: styled HTML, logo embedded, Word letter attached."""
    r = letter["record"]
    msg = EmailMessage()
    msg["To"] = r["email"]
    msg["From"] = ORG_NAME  # a label, not an address - no login needed
    msg["Subject"] = SUBJECT_TEMPLATE.format(name=r["donor_name"], org=ORG_NAME)

    # Plain text first (for mail apps that won't show HTML), then the real thing.
    msg.set_content(letter["body_text"])
    msg.add_alternative(
        f"<!doctype html><html><body>{letter['html_mail']}</body></html>",
        subtype="html",
    )

    # Embed the logo inside the HTML part so it displays in the message itself.
    logo_path = letter.get("logo_path")
    if logo_path and os.path.exists(logo_path):
        html_part = msg.get_payload()[-1]
        subtype = "jpeg" if logo_path.lower().endswith((".jpg", ".jpeg")) else "png"
        with open(logo_path, "rb") as f:
            html_part.add_related(
                f.read(), maintype="image", subtype=subtype, cid="<logoimg>"
            )

    # The formal Word letter rides along as an attachment too.
    with open(letter["letter_path"], "rb") as f:
        msg.add_attachment(
            f.read(),
            maintype="application",
            subtype=_DOCX_SUBTYPE,
            filename=os.path.basename(letter["letter_path"]),
        )

    with open(eml_path, "wb") as f:
        BytesGenerator(f).flatten(msg)


def _write_review(records, prepared, change_log):
    flagged = [r for r in records if r["status"] == "flagged"]
    skipped = [r for r in records if r["status"] == "skipped"]
    out = []
    out.append("# Review before sending\n")
    out.append(
        f"**{len(prepared)} ready, {len(flagged)} flagged, {len(skipped)} skipped "
        "- 0 sent by this tool. Sending is yours.**\n"
    )
    out.append(
        "**To send these, open `SEND.html` in your browser.** Click the button beside "
        "each donor and their email opens already filled in - address, subject, and "
        "letter. Read it, then press Send. No copying and pasting.\n"
    )
    out.append(
        "*(If you use a desktop mail program instead, the `outbox` folder has the same "
        "emails as `.eml` files - double-click one to open it, Word letter attached.)*\n"
    )
    out.append("## Prepared emails\n")
    if prepared:
        out.append("| Donor | Email | Amount | File |")
        out.append("| --- | --- | --- | --- |")
        for item in prepared:
            r = item["record"]
            out.append(
                f"| {r['donor_name']} | {r['email']} | ${float(r['amount']):,.2f} | "
                f"{os.path.basename(item['eml_path'])} |"
            )
    else:
        out.append("_None._")
    out.append("")
    out.append("## Held back (check these by hand)\n")
    held = flagged + skipped
    if held:
        out.append("| Donor | Status | Reason |")
        out.append("| --- | --- | --- |")
        for r in held:
            out.append(
                f"| {r.get('donor_name') or '(no name)'} | {r['status']} | "
                f"{r.get('notes') or ''} |"
            )
    else:
        out.append("_None - every row was ready._")
    out.append("")
    out.append("## What the cleaning step did\n")
    for entry in change_log:
        out.append(f"- {entry}")
    out.append("")
    with open(REVIEW_MD, "w", encoding="utf-8") as f:
        f.write("\n".join(out))


def _subject_for(record):
    return SUBJECT_TEMPLATE.format(name=record["donor_name"], org=ORG_NAME)


def _write_send_page(prepared):
    """Build SEND.html - one click per donor, no copying and pasting.

    Each button opens a Gmail compose window with the address, subject and
    letter already filled in. Gmail still requires YOU to press Send, which is
    exactly the point: the tool never sends anything itself, and never needs
    your password.
    """
    cards = []
    gmail_urls = []
    for index, item in enumerate(prepared):
        record = item["record"]
        subject = _subject_for(record)

        # Address + subject go in the link. The letter itself goes via the
        # clipboard, because a URL can only carry plain text and we want the
        # donor to receive the real, formatted letter with the logo.
        gmail_urls.append(
            "https://mail.google.com/mail/?view=cm&fs=1"
            f"&to={quote(record['email'])}&su={quote(subject)}"
        )

        cards.append(f"""
    <div class="card" id="card{index}">
      <div class="row">
        <input type="checkbox" class="done" id="done{index}" onchange="mark({index})">
        <div class="who">
          <strong>{html_lib.escape(record['donor_name'])}</strong>
          <span class="email">{html_lib.escape(record['email'])}</span>
        </div>
        <div class="amount">${float(record['amount']):,.2f}</div>
        <div class="actions">
          <button class="btn primary" onclick="sendOne({index})">
            Copy letter &amp; open Gmail
          </button>
        </div>
      </div>
      <details id="det{index}">
        <summary>Read the letter</summary>
        <div class="letter" id="letter{index}">{item['html_web']}</div>
      </details>
    </div>""")

    page = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Send your thank-you emails</title>
<style>
  :root {{ color-scheme: light dark; }}
  body {{ font-family: system-ui, -apple-system, "Segoe UI", sans-serif;
         max-width: 900px; margin: 0 auto; padding: 24px; line-height: 1.5; }}
  h1 {{ margin-bottom: 4px; }}
  .lede {{ color: #666; margin-top: 0; }}
  .banner {{ background: #fff8e1; border: 1px solid #ffe082; color: #5d4200;
            padding: 12px 16px; border-radius: 8px; margin: 16px 0; }}
  @media (prefers-color-scheme: dark) {{
    .lede, .email {{ color: #aaa; }}
    .banner {{ background: #2e2a1a; border-color: #6b5a1f; color: #f0e2b6; }}
    .card {{ border-color: #444; }}
    .btn {{ border-color: #555; }}
  }}
  .counter {{ font-weight: 600; margin: 16px 0; }}
  .card {{ border: 1px solid #ddd; border-radius: 10px; padding: 12px 16px;
          margin-bottom: 10px; }}
  .card.sent {{ opacity: .45; }}
  .row {{ display: flex; align-items: center; gap: 12px; flex-wrap: wrap; }}
  .who {{ display: flex; flex-direction: column; flex: 1 1 220px; }}
  .email {{ color: #666; font-size: .9em; }}
  .amount {{ font-variant-numeric: tabular-nums; font-weight: 600; }}
  .actions {{ display: flex; gap: 8px; }}
  .btn {{ display: inline-block; padding: 8px 14px; border-radius: 6px;
         border: 1px solid #ccc; text-decoration: none; color: inherit;
         white-space: nowrap; font: inherit; cursor: pointer;
         background: transparent; }}
  .btn.primary {{ background: #1a73e8; border-color: #1a73e8; color: #fff; }}
  .btn:hover {{ filter: brightness(1.08); }}
  .banner ol {{ margin: 8px 0 0; padding-left: 22px; }}
  .banner li {{ margin-bottom: 4px; }}
  details {{ margin-top: 8px; }}
  summary {{ cursor: pointer; color: #1a73e8; }}
  /* The letter preview is also what gets copied, so show it as it will look. */
  .letter {{ background: #fff; color: #202124; border: 1px solid #ddd;
            border-radius: 6px; padding: 20px 24px; margin-top: 8px;
            overflow-x: auto; }}
  .letter img {{ max-width: 200px; height: auto; }}
  .toast {{ position: fixed; left: 50%; bottom: 24px; transform: translateX(-50%);
           background: #202124; color: #fff; padding: 12px 20px;
           border-radius: 8px; opacity: 0; pointer-events: none;
           transition: opacity .2s; max-width: 90vw; text-align: center; }}
  .toast.show {{ opacity: 1; }}
</style>
</head>
<body>
<h1>Send your thank-you emails</h1>
<p class="lede">{len(prepared)} letter(s) ready. Nothing has been sent.</p>

<div class="banner">
  <strong>For each donor, do these three things:</strong>
  <ol>
    <li>Click the blue <strong>Copy letter &amp; open Gmail</strong> button.
        Gmail opens with their address and subject already filled in.</li>
    <li>Click inside the big empty message box and press
        <strong>Ctrl&nbsp;+&nbsp;V</strong> to paste the letter in. It keeps the
        logo, fonts, bold and colours.</li>
    <li>Read it, then press <strong>Send</strong> in Gmail.</li>
  </ol>
  This page never sends anything and never sees your password.
</div>

<p class="counter"><span id="count">0</span> of {len(prepared)} marked sent</p>
{"".join(cards)}

<div class="toast" id="toast">Letter copied - now press Ctrl+V in Gmail</div>

<script>
  const TOTAL = {len(prepared)};
  const GMAIL = {json.dumps(gmail_urls)};

  function key(i) {{ return "sent_" + location.pathname + "_" + i; }}

  function refresh() {{
    let n = 0;
    for (let i = 0; i < TOTAL; i++) {{
      const box = document.getElementById("done" + i);
      if (box && box.checked) n++;
    }}
    document.getElementById("count").textContent = n;
  }}
  function mark(i) {{
    const box = document.getElementById("done" + i);
    document.getElementById("card" + i).classList.toggle("sent", box.checked);
    try {{ localStorage.setItem(key(i), box.checked ? "1" : ""); }} catch (e) {{}}
    refresh();
  }}
  function tick(i) {{
    const box = document.getElementById("done" + i);
    if (box && !box.checked) {{ box.checked = true; mark(i); }}
  }}

  function toast(message) {{
    const el = document.getElementById("toast");
    el.textContent = message;
    el.classList.add("show");
    setTimeout(() => el.classList.remove("show"), 3500);
  }}

  // Copy the letter as RICH text, so the logo, fonts, bold and colours all
  // survive the paste into Gmail. Plain text is included as a fallback for
  // anything that can't take HTML.
  async function copyLetter(i) {{
    const el = document.getElementById("letter" + i);
    const html = el.innerHTML;
    const text = el.innerText;
    try {{
      await navigator.clipboard.write([
        new ClipboardItem({{
          "text/html": new Blob([html], {{ type: "text/html" }}),
          "text/plain": new Blob([text], {{ type: "text/plain" }})
        }})
      ]);
      return true;
    }} catch (e) {{
      // Older browsers: select the letter and use the classic copy command.
      const details = document.getElementById("det" + i);
      const wasOpen = details.open;
      details.open = true;
      const range = document.createRange();
      range.selectNodeContents(el);
      const sel = window.getSelection();
      sel.removeAllRanges();
      sel.addRange(range);
      let ok = false;
      try {{ ok = document.execCommand("copy"); }} catch (err) {{ ok = false; }}
      sel.removeAllRanges();
      details.open = wasOpen;
      return ok;
    }}
  }}

  async function sendOne(i) {{
    const copied = await copyLetter(i);
    window.open(GMAIL[i], "_blank", "noopener");
    tick(i);
    toast(copied
      ? "Letter copied - click in the Gmail message box and press Ctrl+V"
      : "Could not copy automatically - open 'Read the letter', select it, and copy");
  }}

  for (let i = 0; i < TOTAL; i++) {{
    try {{
      if (localStorage.getItem(key(i))) {{
        document.getElementById("done" + i).checked = true;
        document.getElementById("card" + i).classList.add("sent");
      }}
    }} catch (e) {{}}
  }}
  refresh();
</script>
</body>
</html>
"""
    with open(SEND_HTML, "w", encoding="utf-8") as handle:
        handle.write(page)


def stage_prepare(records, letters, change_log):
    os.makedirs(OUTBOX_DIR, exist_ok=True)
    prepared = []
    for letter in letters:
        stem = os.path.splitext(os.path.basename(letter["letter_path"]))[0]
        eml_path = os.path.join(OUTBOX_DIR, stem + ".eml")
        _write_eml(letter, eml_path)
        letter["eml_path"] = eml_path
        prepared.append(letter)
    _write_review(records, prepared, change_log)
    _write_send_page(prepared)
    return prepared


# ==========================================================================
#  Putting it together
# ==========================================================================
def _write_cleaned_csv(records):
    fields = ["donor_name", "email", "amount", "date", "donor_type", "status", "notes"]
    with open(CLEANED_CSV, "w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for r in records:
            writer.writerow({k: r.get(k, "") for k in fields})


def _in_colab():
    try:
        import google.colab  # noqa: F401
        return True
    except Exception:
        return False


def _gather_inputs():
    """Return (xlsx_path, docx_path), creating blank starter files if needed."""
    xlsx_path, docx_path = "donations.xlsx", "letter_template.docx"
    if _in_colab():
        try:
            from google.colab import files
            print(
                "Optional: upload YOUR donations spreadsheet (.xlsx) and letter "
                "template (.docx) now.\nOr choose nothing to get blank starter "
                "files you can fill in."
            )
            uploaded = files.upload()
            for name in uploaded:
                low = name.lower()
                if low.endswith(".xlsx"):
                    xlsx_path = name
                elif low.endswith(".docx"):
                    docx_path = name
        except Exception as exc:
            print(f"(Skipping upload: {exc})")
    if not os.path.exists(xlsx_path):
        generate_donations(xlsx_path)
        print(f"Created a blank donation sheet for you: {xlsx_path}")
    if not os.path.exists(docx_path):
        generate_template(docx_path)
        print(f"Created a letter template for you: {docx_path}")
    return xlsx_path, docx_path


def _explain_locked_file(exc):
    """Turn a raw PermissionError traceback into something a human can act on.

    On Windows this nearly always means the file is still open in Excel or Word.
    """
    name = getattr(exc, "filename", None) or "one of your files"
    print("\n" + "-" * 60)
    print(f"Windows won't let the tool open:  {name}")
    print("\nThis almost always means the file is still open in another program.")
    print("\n  1. Close it in Excel / Word  (save your changes first).")
    print("  2. Close the File Explorer preview pane if it's showing the file.")
    print("  3. If it's in OneDrive, wait for the sync icon to stop spinning.")
    print("  4. Then press Run again.")
    print("-" * 60)


def main():
    try:
        _run()
    except PermissionError as exc:
        _explain_locked_file(exc)


def _run():
    print("Donation Thank-You Tool")
    print("=" * 40)
    if USE_CLAUDE_AI:
        print("Mode: Claude AI (uses your paid API key)")
    else:
        print("Mode: FREE - no key, no internet, no cost. Nothing leaves this computer.")
    print("Steps: 1) clean spreadsheet  2) write letters  3) prepare emails (you send)\n")

    xlsx_path, docx_path = _gather_inputs()

    # If the sheet is still just the "Insert ... here" boxes, say so and stop.
    _header, _rows = read_xlsx(xlsx_path)
    if looks_unfilled(_rows):
        print("\n" + "-" * 60)
        print("Your donation sheet is empty - nothing to thank anyone for yet.")
        print(f"\n  1. Open this file:  {os.path.abspath(xlsx_path)}")
        print("  2. Type over the three boxes:")
        print("       Donor Name  |  Donation Amount  |  Email")
        print("       (e.g.  Jane Doe  |  50  |  jane@example.org )")
        print("  3. Add one row per additional donor.")
        print("  4. Save the file, then press Run again.")
        print("-" * 60)
        return

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("\n[1/3] Cleaning the donation log...")
    records, change_log = stage_clean(xlsx_path)
    _write_cleaned_csv(records)
    for entry in change_log:
        print(f"    - {entry}")

    print("\n[2/3] Writing personalised letters...")
    letters = stage_personalize(records, docx_path)
    print(f"    {len(letters)} letter(s) written")

    print("\n[3/3] Preparing the outbox...")
    prepared = stage_prepare(records, letters, change_log)

    ready = sum(1 for r in records if r["status"] == "ready")
    flagged = sum(1 for r in records if r["status"] == "flagged")
    skipped = sum(1 for r in records if r["status"] == "skipped")
    print("\nDone.")
    print(f"  {ready} ready, {flagged} flagged, {skipped} skipped.")
    print(f"  {len(prepared)} email(s) prepared - 0 sent by this tool. Sending is yours.")
    print("\n  >> TO SEND: open  output/SEND.html  in your browser.")
    print("     For each donor: click the blue button (Gmail opens, address filled in),")
    print("     press Ctrl+V to paste the letter, then press Send.")
    print("     The paste keeps your logo, fonts, bold and colours.")
    print("\n  A summary of what was held back is in output/REVIEW.md.")

    # Pop the send page open so there's nothing left to go hunting for.
    if OPEN_SEND_PAGE and prepared and not _in_colab():
        try:
            webbrowser.open("file://" + os.path.abspath(SEND_HTML).replace("\\", "/"))
        except Exception:
            pass  # not worth failing a good run over

    # In Colab there's no local folder to browse, so zip it into one download.
    if _in_colab():
        zip_name = "donation_output.zip"
        with zipfile.ZipFile(zip_name, "w", zipfile.ZIP_DEFLATED) as zf:
            for root, _dirs, filenames in os.walk(OUTPUT_DIR):
                for filename in filenames:
                    full = os.path.join(root, filename)
                    zf.write(full, os.path.relpath(full, "."))
        try:
            from google.colab import files
            print(f"\nDownloading {zip_name} - open it and start with REVIEW.md.")
            files.download(zip_name)
        except Exception as exc:
            print(f"(Could not auto-download: {exc}). Find {zip_name} in the file panel.")


if __name__ == "__main__":
    main()
